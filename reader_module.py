# Required imports
import os
import json
import pandas as pd
import numpy as np
from PIL import Image
import docx
import pdfplumber
from bs4 import BeautifulSoup
import librosa
from io import BytesIO
import zipfile

class DataExpander:
    def __init__(self):
        pass

    def expand(self, file_path, modality: str = None):
        try:
            file_extension = file_path.lower().split('.')[-1]

            if file_extension in ('jpg', 'jpeg', 'png', 'jpe'):
                return self._read_image(file_path)

            elif file_extension == 'docx':
                return self._read_docx(file_path)

            elif file_extension == 'csv':
                return self._read_csv(file_path)

            elif file_extension == 'pdf':
                return self._read_pdf(file_path)

            elif file_extension == 'txt':
                return self._read_txt(file_path)

            elif file_extension == 'html':
                return self._read_html(file_path)

            elif file_extension == 'json':
                return self._read_json(file_path)

            elif file_extension in ('xls', 'xlsx'):
                return self._read_xlsx(file_path)

            elif file_extension == 'mp3':
                return self._read_mp3(file_path)

            else:
                raise ValueError(f"Unsupported file type: {file_extension}. Supported: PDF, DOCX, XLSX, IMAGE, TXT, HTML, CSV, JSON, MP3")

        except FileNotFoundError:
            raise FileNotFoundError(f"File not found: {file_path}")
        except Exception as e:
            raise ValueError(f"Error processing file: {str(e)}")

    def _read_pdf(self, file_path):
        try:
            text = []
            images = []
            tables = []

            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text.extend([line.strip() for line in page_text.splitlines() if line.strip()])

                    # Extract tables
                    table_data = page.extract_tables()
                    for tbl in table_data:
                        if tbl:
                            tables.append(pd.DataFrame(tbl))

                    # Extract image
                    im = page.to_image(resolution=150)
                    images.append(im.original)

            return {
                "type": ["text"] + (["images"] if images else []) + (["tables"] if tables else []),
                "content": {
                    "text": text,  # <-- list of lines
                    "images": images,
                    "tables": tables
                },
                "meta": {
                    "filename": os.path.basename(file_path),
                    "format": "pdf",
                    "pages": len(pdf.pages),
                    "num_tables": len(tables),
                    "word_count": sum(len(p.split()) for p in text)
                }
            }

        except Exception as e:
            return self._error_response(file_path, e)
        
    def _read_docx(self, file_path):
        try:
            doc = docx.Document(file_path)
            full_text = []
            tables = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

            # Extract tables
            for table in doc.tables:
                data = []
                for row in table.rows:
                    data.append([cell.text.strip() for cell in row.cells])
                tables.append(pd.DataFrame(data))

            # Extract images
            images = []
            with zipfile.ZipFile(file_path) as docx_zip:
                for file in docx_zip.namelist():
                    if file.startswith("word/media/") and file.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".gif")):
                        with docx_zip.open(file) as image_file:
                            image = Image.open(BytesIO(image_file.read()))
                            images.append(image.copy())

            return {
                "type": ["text"] + (["tables"] if tables else []) + (["images"] if images else []),
                "content": {
                    "text": full_text,  # <-- list of paragraphs
                    "images": images,
                    "tables": tables
                },
                "meta": {
                    "filename": os.path.basename(file_path),
                    "format": "docx",
                    "num_paragraphs": len(full_text),
                    "num_tables": len(tables),
                    "num_images": len(images),
                    "word_count": sum(len(p.split()) for p in full_text)
                }
            }

        except Exception as e:
            return self._error_response(file_path, e)
        
    # def _read_txt(self, file_path):
    #     try:
    #         with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
    #             lines = [line.strip() for line in f if line.strip()]
    #         return {
    #             "type": "text",
    #             "content": lines,  # <-- list of non-empty lines
    #             "meta": {
    #                 "original_filename": os.path.basename(file_path),
    #                 "original_extension": os.path.splitext(file_path)[1].lower(),
    #                 "num_lines": len(lines),
    #                 "word_count": sum(len(line.split()) for line in lines)
    #             }
    #         }
    #     except Exception as e:
    #         return self._error_response(file_path, e)

    def _read_txt(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            return {
                "type": "text",
                "content": text,  # <-- full text as a single string
                "meta": {
                    "original_filename": os.path.basename(file_path),
                    "original_extension": os.path.splitext(file_path)[1].lower(),
                    "num_lines": text.count('\n') + 1,
                    "word_count": len(text.split())
                }
            }
        except Exception as e:
            return self._error_response(file_path, e)

    def _read_xlsx(self, file_path, sheet_name=0):
        try:
            xls = pd.ExcelFile(file_path)
            if isinstance(sheet_name, int):
                if sheet_name < 0 or sheet_name >= len(xls.sheet_names):
                    raise ValueError(f"Invalid sheet index {sheet_name}")
                active_sheet_name = xls.sheet_names[sheet_name]
            elif isinstance(sheet_name, str):
                if sheet_name not in xls.sheet_names:
                    raise ValueError(f"Sheet '{sheet_name}' not found")
                active_sheet_name = sheet_name
            else:
                raise TypeError("Sheet name must be int or str")

            df = pd.read_excel(file_path, sheet_name=active_sheet_name)

            return {
                "type": "table",
                "content": df,
                "meta": {
                    "original_filename": os.path.basename(file_path),
                    "original_extension": os.path.splitext(file_path)[1].lower(),
                    "shape": df.shape,
                    "columns": df.columns.tolist(),
                    "sheet_names": xls.sheet_names,
                    "active_sheet_name": active_sheet_name
                }
            }
        except Exception as e:
            return self._error_response(file_path, e)

    def _read_image(self, file_path):
        try:
            img = Image.open(file_path).convert('RGB')
            img_array = np.array(img)
            return {
                "type": "image",
                "content": img_array,
                "meta": {
                    "original_filename": os.path.basename(file_path),
                    "original_extension": os.path.splitext(file_path)[1].lower(),
                    "size_pixels": img.size,  # (width, height)
                    "mode": img.mode
                }
            }
        except Exception as e:
            return self._error_response(file_path, e)

    def _read_html(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                soup = BeautifulSoup(f, 'lxml')
            text_content = soup.get_text(separator='\n', strip=True)
            title_tag = soup.find('title')
            title = title_tag.string.strip() if title_tag and title_tag.string else "N/A"
            links = soup.find_all('a')

            return {
                "type": "text",
                "content": text_content,
                "meta": {
                    "original_filename": os.path.basename(file_path),
                    "original_extension": os.path.splitext(file_path)[1].lower(),
                    "title": title,
                    "num_links": len(links)
                }
            }
        except Exception as e:
            return self._error_response(file_path, e)

    def _read_csv(self, file_path, delimiter=','):
        try:
            df = pd.read_csv(file_path, delimiter=delimiter)
            return {
                "type": "table",
                "content": df,
                "meta": {
                    "original_filename": os.path.basename(file_path),
                    "original_extension": os.path.splitext(file_path)[1].lower(),
                    "shape": df.shape,
                    "columns": df.columns.tolist(),
                    "delimiter": delimiter
                }
            }
        except Exception as e:
            return self._error_response(file_path, e)

    def _read_json(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            meta_info = {
                "original_filename": os.path.basename(file_path),
                "original_extension": os.path.splitext(file_path)[1].lower(),
            }
            if isinstance(data, list):
                meta_info["is_list"] = True
                meta_info["num_elements"] = len(data)
            elif isinstance(data, dict):
                meta_info["is_list"] = False
                meta_info["num_top_level_keys"] = len(data.keys())

            return {
                "type": "table",
                "content": pd.DataFrame(data),
                "meta": meta_info
            }
        except Exception as e:
            return self._error_response(file_path, e)

    def _read_mp3(self, file_path):
        try:
            y, sr = librosa.load(file_path)
            return {
                "type": "audio",
                "content": y,
                "meta": {
                    "original_filename": os.path.basename(file_path),
                    "original_extension": os.path.splitext(file_path)[1].lower(),
                    "sample_rate": sr,
                    "duration_sec": len(y) / sr
                }
            }
        except Exception as e:
            return self._error_response(file_path, e)

    def _error_response(self, file_path, error):
        return {
            "type": "error",
            "content": str(error),
            "meta": {
                "original_filename": os.path.basename(file_path),
                "original_extension": os.path.splitext(file_path)[1].lower()
            }
        }
