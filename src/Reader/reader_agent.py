import os
import json
import pandas as pd
import numpy as np
import docx
import pdfplumber
from bs4 import BeautifulSoup
import librosa 
import cv2
import io
import fitz

class DataExpander:
    """
    Một lớp để trích xuất và chuyển đổi nội dung từ nhiều định dạng tệp khác nhau.
    Nội dung hình ảnh và âm thanh được trả về dưới dạng np.array.
    Nội dung bảng biểu được trả về dưới dạng pd.DataFrame.
    Tất cả các phương thức _read trả về một danh sách các từ điển.
    """
    def __init__(self):
        pass
    
    def expand(self, file_path, modality: str = None):
        """
        Phương thức chính để đọc một tệp và trích xuất nội dung của nó.
        Trả về một danh sách các từ điển, mỗi từ điển đại diện cho một phần dữ liệu.
        """
        try:
            file_extension = file_path.lower().split('.')[-1]
                
            if file_extension in ('jpg', 'jpeg', 'png', 'jpe'):
                return self._read_image(file_path)
            
            elif file_extension == 'docx':
                return self._read_docx(file_path)
            
            elif file_extension == 'csv':
                return self._read_csv(file_path)
            
            elif file_extension == 'pdf':
                return self._read_pdf(file_path)
            
            elif file_extension == 'txt':
                return self._read_txt(file_path)
            
            elif file_extension == 'html':
                return self._read_html(file_path)
            
            elif file_extension == 'json':
                return self._read_json(file_path)
            
            elif file_extension in ('xls', 'xlsx'):
                return self._read_xlsx(file_path)
            
            elif file_extension in ('mp3'):
                return self._read_audio(file_path)
            
            else:
                raise ValueError(f"Loại tệp không được hỗ trợ: {file_extension}")
                        
        except FileNotFoundError:
            raise FileNotFoundError(f"Không tìm thấy tệp: {file_path}")
        except Exception as e:
            raise ValueError(f"Lỗi khi xử lý tệp: {str(e)}")

    def _create_error_output(self, file_path, e):
        """Tạo một danh sách chứa một từ điển thông báo lỗi chuẩn."""
        return [{
            "type": "error",
            "content": str(e),
            "meta": {
                "original_filename": os.path.basename(file_path),
                "original_extension": os.path.splitext(file_path)[1].lower()
            }
        }]

    def _read_pdf(self, file_path):
        """
        Trích xuất văn bản, bảng và hình ảnh từ PDF.
        Trả về danh sách các từ điển với type là 'text', 'table', hoặc 'image'.
        """
        try:
            results = []
            
            # 1. Trích xuất văn bản và bảng bằng pdfplumber
            with pdfplumber.open(file_path) as pdf:
                full_text_content = ""
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        full_text_content += page_text + "\n"
                    
                    for table_data in page.extract_tables():
                        if table_data:
                            df = pd.DataFrame(table_data[1:], columns=table_data[0])
                            results.append({
                                "type": "table",
                                "content": df,
                                "meta": {
                                    "original_filename": os.path.basename(file_path),
                                    "original_extension": ".pdf",
                                    "source": f"page_{i+1}"
                                }
                            })
                
                if full_text_content.strip():
                    results.append({
                        "type": "text",
                        "content": full_text_content.strip(),
                        "meta": {
                            "original_filename": os.path.basename(file_path),
                            "original_extension": ".pdf"
                        }
                    })
            
            # 2. Trích xuất hình ảnh bằng PyMuPDF
            with fitz.open(file_path) as doc:
                for page_index, page in enumerate(doc):
                    for img in page.get_images(full=True):
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        
                        nparr = np.frombuffer(image_bytes, np.uint8)
                        img_np = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                        if img_np is not None:
                            img_np_rgb = cv2.cvtColor(img_np, cv2.COLOR_BGR2RGB)
                            results.append({
                                "type": "image",
                                "content": img_np_rgb,
                                "meta": {
                                    "original_filename": os.path.basename(file_path),
                                    "original_extension": ".pdf",
                                    "source": f"page_{page_index+1}",
                                    "size_pixels": img_np_rgb.shape
                                }
                            })
            
            return results if results else self._create_error_output(file_path, "Không tìm thấy nội dung trong PDF")

        except Exception as e:
            return self._create_error_output(file_path, e)

    def _read_docx(self, file_path):
        """Trích xuất văn bản, bảng và hình ảnh từ DOCX. Trả về danh sách các từ điển với type là 'text', 'table', hoặc 'image'."""
        try:
            results = []
            doc = docx.Document(file_path)
            
            # Trích xuất văn bản
            text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
            if text_parts:
                full_text_content = "\n".join(text_parts)
                results.append({
                    "type": "text",
                    "content": full_text_content,
                    "meta": {
                        "original_filename": os.path.basename(file_path),
                        "original_extension": ".docx"
                    }
                })
            
            # Trích xuất bảng
            for i, table in enumerate(doc.tables):
                data = [[cell.text for cell in row.cells] for row in table.rows]
                df = pd.DataFrame(data[1:], columns=data[0])
                results.append({
                    "type": "table",
                    "content": df,
                    "meta": {
                        "original_filename": os.path.basename(file_path),
                        "original_extension": ".docx",
                        "source": f"table_{i+1}"
                    }
                })

            # Trích xuất hình ảnh
            image_parts = [part for part in doc.part.related_parts.values() if "image" in part.content_type]
            for i, part in enumerate(image_parts):
                image_bytes = part.blob
                nparr = np.frombuffer(image_bytes, np.uint8)
                img_np = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                if img_np is not None:
                    img_np_rgb = cv2.cvtColor(img_np, cv2.COLOR_BGR2RGB)
                    results.append({
                        "type": "image",
                        "content": img_np_rgb,
                        "meta": {
                            "original_filename": os.path.basename(file_path),
                            "original_extension": ".docx",
                            "source": f"image_{i+1}",
                            "size_pixels": img_np_rgb.shape
                        }
                    })
            
            return results if results else self._create_error_output(file_path, "Không tìm thấy nội dung trong DOCX")

        except Exception as e:
            return self._create_error_output(file_path, e)

    def _read_html(self, file_path):
        """Trích xuất văn bản và bảng (DataFrame) từ tệp HTML. Trả về danh sách các từ điển."""
        results = []
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            soup = BeautifulSoup(content, 'lxml')
            text_content = soup.get_text(separator='\n', strip=True)
            if text_content:
                results.append({
                    "type": "text",
                    "content": text_content,
                    "meta": { "original_filename": os.path.basename(file_path) }
                })
            try:
                tables = pd.read_html(io.StringIO(content))
                for i, df in enumerate(tables):
                    results.append({
                        "type": "table",
                        "content": df,
                        "meta": {
                            "original_filename": os.path.basename(file_path),
                            "source": f"table_{i+1}"
                        }
                    })
            except ValueError:
                pass
            return results if results else self._create_error_output(file_path, "Không tìm thấy nội dung HTML hợp lệ")
        except Exception as e:
            return self._create_error_output(file_path, e)
            
    def _read_xlsx(self, file_path, sheet_name=None):
        """Đọc tất cả các trang tính từ tệp Excel. Trả về danh sách các từ điển chứa DataFrame."""
        results = []
        try:
            xls = pd.ExcelFile(file_path)
            sheet_names_to_read = xls.sheet_names if sheet_name is None else [sheet_name]
            for s_name in sheet_names_to_read:
                df = pd.read_excel(xls, sheet_name=s_name)
                results.append({
                    "type": "table",
                    "content": df,
                    "meta": {
                        "original_filename": os.path.basename(file_path),
                        "sheet_name": s_name
                    }
                })
            return results
        except Exception as e:
            return self._create_error_output(file_path, e)

    def _read_image(self, file_path):
        """Đọc tệp hình ảnh và trả về danh sách chứa một từ điển với np.array."""
        try:
            img = cv2.imread(file_path)
            if img is None:
                raise ValueError("Không thể đọc hình ảnh")
            img_array = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
            return [{
                "type": "image",
                "content": img_array,
                "meta": {
                    "original_filename": os.path.basename(file_path),
                    "size_pixels": img_array.shape
                }
            }]
        except Exception as e:
            return self._create_error_output(file_path, e)

    def _read_txt(self, file_path):
        """Đọc tệp văn bản và trả về danh sách chứa một từ điển."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f: 
                content = f.read()
            return [{
                "type": "text",
                "content": content,
                "meta": { 
                    "original_filename": os.path.basename(file_path) 
                }
            }]
        except Exception as e:
            return self._create_error_output(file_path, e)

    def _read_csv(self, file_path, delimiter=','):
        """Đọc tệp CSV và trả về danh sách chứa một từ điển với DataFrame."""
        try:
            df = pd.read_csv(file_path, delimiter=delimiter)
            return [{
                "type": "table",
                "content": df,
                "meta": { 
                    "original_filename": os.path.basename(file_path) 
                }
            }]
        except Exception as e:
            return self._create_error_output(file_path, e)

    def _read_json(self, file_path):
        """Đọc tệp JSON và trả về danh sách chứa một từ điển với pd.DataFrame."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return [{
                "type": "table",
                "content": pd.DataFrame(data),
                "meta": { 
                    "original_filename": os.path.basename(file_path) 
                }
            }]
        except Exception as e:
            return self._create_error_output(file_path, e)

    def _read_audio(self, file_path):
        """Đọc tệp âm thanh và trả về danh sách chứa một từ điển với np.array."""
        try:
            y, sr = librosa.load(file_path, sr=None)
            return [{
                "type": "audio",
                "content": y,
                "meta": {
                    "original_filename": os.path.basename(file_path),
                    "sampling_rate": sr
                }
            }]
        except Exception as e:
            return self._create_error_output(file_path, e)
